import os
import time
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import streamlit as st
from pipeline import E2EQGPipeline
import tempfile

def save_to_docx(questions):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc = Document()
        doc.add_heading('Questions', level=1)
        for i, question in enumerate(questions, 1):
            doc.add_paragraph(f"Q{i}: {question}")
        doc.save(tmp.name)
        return tmp.name

def save_to_pdf(questions):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        c = canvas.Canvas(tmp.name, pagesize=letter)
        c.drawString(72, 800, 'Questions')
        y = 780
        for i, question in enumerate(questions, 1):
            c.drawString(72, y - i * 20, f"Q{i}: {question}")
        c.save()
        return tmp.name

def main():
    st.title("End-to-End Question Generation")
    st.write("Generate questions from your text using a pretrained NLP model.")
    
    context = st.text_area("Enter your text here:")
    
    if st.button("Generate Questions"):
        with st.spinner('Generating questions...'):
            start_time = time.time()
            nlp = E2EQGPipeline(model_name="valhalla/t5-base-e2e-qg", tokenizer_name="valhalla/t5-base-e2e-qg", use_cuda=False)
            questions = nlp(context, num_questions=10)  # Always generate 10 questions
            end_time = time.time()
            
            if not questions:
                st.warning("No questions generated. Please try with a different text.")
                return

            unique_questions = list(set(questions))

            # If fewer unique questions are generated, use the unique ones available
            if len(unique_questions) < 10:
                st.warning(f"Only {len(unique_questions)} unique questions generated.")
                num_questions = len(unique_questions)
            else:
                num_questions = 10
            
            doc_filename = save_to_docx(unique_questions[:num_questions])
            pdf_filename = save_to_pdf(unique_questions[:num_questions])

            total_time = end_time - start_time

            st.success("Questions generated successfully!")
            st.write(f"### CPU: {total_time:.4f} seconds")
            
            st.write("### Questions")
            for i, question in enumerate(unique_questions[:num_questions], 1):
                st.write(f"Q{i}: {question}")

            st.write("### Download Links")
            with open(doc_filename, "rb") as file:
                st.download_button(label="Download Word Document", data=file, file_name="questions_generated.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            with open(pdf_filename, "rb") as file:
                st.download_button(label="Download PDF Document", data=file, file_name="questions_generated.pdf", mime="application/pdf")

if __name__ == "_main_":
    main()